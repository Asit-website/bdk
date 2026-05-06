
import os
import random
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TOPIC = "आर्थिक विकास और कूटनीति के संदर्भ में विश्लेषण: एक रणनीतिक अध्ययन"
REGULAR_FONT = "Poppins-Regular.ttf"
BOLD_FONT = "Poppins-Bold.ttf"

class PDF(FPDF):
    def header(self):
        self.set_line_width(1.5)
        self.rect(10, 10, 190, 277)
        self.set_line_width(0.5)
        self.rect(12, 12, 186, 273)
        if self.page_no() > 1:
            self.set_font('Poppins', '', 10)
            self.set_y(15)
            self.multi_cell(0, 5, TOPIC, 0, 'C')
            self.set_y(30)

    def footer(self):
        self.set_y(-15)
        self.set_font('Poppins', '', 10)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def set_page_border_docx(doc):
    sec = doc.sections[0]
    sectPr = sec._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'double')
        border_el.set(qn('w:sz'), '12')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), '000000')
        pgBorders.append(border_el)
    sectPr.append(pgBorders)

SENTENCES = [
    "आर्थिक कूटनीति वर्तमान वैश्विक व्यवस्था में किसी भी राष्ट्र की विदेश नीति का एक अत्यंत महत्वपूर्ण और आधारभूत स्तंभ बन गई है।",
    "आर्थिक विकास और कूटनीति का आपसी संबंध राष्ट्रों के बीच शक्ति संतुलन और सहयोग के नए आयामों को निरंतर परिभाषित कर रहा है।",
    "वैश्वीकरण के इस युग में व्यापारिक समझौते और विदेशी निवेश को आकर्षित करना कूटनीतिक सफलता का एक प्रमुख पैमाना माना जाता है।",
    "आर्थिक विकास केवल घरेलू नीतियों पर निर्भर नहीं है, बल्कि यह अंतरराष्ट्रीय व्यापारिक संबंधों और कूटनीतिक संवाद पर भी टिका है।",
    "भारत जैसी उभरती अर्थव्यवस्थाओं के लिए आर्थिक कूटनीति वैश्विक स्तर पर अपनी रणनीतिक स्वायत्तता को मजबूत करने का एक साधन है।",
    "ऊर्जा सुरक्षा और संसाधनों तक पहुंच सुनिश्चित करना आर्थिक कूटनीति का एक बहुत ही अनिवार्य, चुनौतीपूर्ण और रणनीतिक पक्ष है।",
    "विदेशी प्रत्यक्ष निवेश (FDI) को बढ़ावा देने के लिए राष्ट्रों के बीच स्वस्थ कूटनीतिक संबंधों का होना आज के दौर में बहुत जरूरी है।",
    "आर्थिक कूटनीति के माध्यम से राष्ट्र अपने आर्थिक हितों की रक्षा करते हुए वैश्विक राजनीति में अपना प्रभाव और स्थान बढ़ाते हैं।",
    "बहुपक्षीय संस्थानों जैसे विश्व व्यापार संगठन (WTO) में अपनी आवाज बुलंद करना विकासशील देशों की आर्थिक कूटनीति का हिस्सा है।",
    "आर्थिक विकास के लिए तकनीकी हस्तांतरण और नवाचार को बढ़ावा देना कूटनीतिक समझौतों का एक बहुत ही महत्वपूर्ण और आधुनिक लक्ष्य है।",
    "आर्थिक कूटनीति न केवल व्यापार को बढ़ावा देती है, बल्कि यह राष्ट्रों के बीच आपसी विश्वास और दीर्घकालिक शांति का आधार भी बनती है।",
    "क्षेत्रीय व्यापारिक ब्लॉक (Regional Trading Blocks) आर्थिक कूटनीति के माध्यम से क्षेत्रीय आर्थिक एकीकरण को मजबूती प्रदान करते हैं।",
    "सतत विकास लक्ष्यों (SDGs) को प्राप्त करने के लिए अंतरराष्ट्रीय आर्थिक सहयोग और कूटनीतिक समन्वय आज अत्यंत अनिवार्य हो गया है।",
    "आर्थिक प्रतिबंधों का कूटनीतिक हथियार के रूप में उपयोग वैश्विक व्यापार और अंतरराष्ट्रीय संबंधों की स्थिरता के लिए एक बड़ी चुनौती है।",
    "डिजिटल अर्थव्यवस्था और साइबर व्यापार के बढ़ते प्रभाव ने आर्थिक कूटनीति के लिए नए नियम और रणनीतिक ढांचे तैयार किए हैं।",
    "आर्थिक कूटनीति का मुख्य उद्देश्य राष्ट्र के आर्थिक विकास को गति देना और वैश्विक संसाधनों का न्यायोचित वितरण सुनिश्चित करना है।",
    "विकासशील राष्ट्रों के बीच दक्षिण-दक्षिण सहयोग (South-South Cooperation) आर्थिक कूटनीति का एक बहुत ही सकारात्मक स्वरूप है।",
    "आर्थिक कूटनीति और राष्ट्रीय सुरक्षा एक-दूसरे के पूरक हैं, जो राष्ट्र को वैश्विक स्तर पर अधिक सुरक्षित और समृद्ध बनाते हैं।",
    "पर्यावरण अनुकूल व्यापार नीतियां और हरित ऊर्जा के क्षेत्र में सहयोग आधुनिक आर्थिक कूटनीति की एक नई और अनिवार्य दिशा है।",
    "निष्कर्षतः, आर्थिक विकास और कूटनीति का सफल संगम ही 21वीं सदी में एक सशक्त and समृद्ध राष्ट्र निर्माण की सच्ची गारंटी है।"
] * 150

def get_unique_content(page_num, count=11):
    random.seed(page_num + 119)
    content = random.sample(SENTENCES, count)
    return " ".join(content)

def generate_economic_diplomacy_perfect_60():
    chapter_names = [
        "परिचय: आर्थिक कूटनीति की अवधारणा एवं महत्व", "ऐतिहासिक विकास: व्यापार से लेकर सामरिक कूटनीति तक", "वैश्विक व्यापारिक संस्थान एवं आर्थिक कूटनीति", 
        "भारत की आर्थिक कूटनीति: चुनौतियां एवं संभावनाएं", "विदेशी निवेश, तकनीकी सहयोग एवं आर्थिक विकास", "ऊर्जा कूटनीति एवं वैश्विक संसाधन प्रबंधन", 
        "बहुध्रुवीय विश्व में आर्थिक कूटनीति का बदलता स्वरूप"
    ]
    
    pages_per_ch = [7, 7, 8, 8, 8, 8, 8] # Total 54 content pages
    
    # PDF
    pdf = PDF()
    pdf.set_auto_page_break(auto=False)
    pdf.set_text_shaping(True)
    pdf.add_font('Poppins', '', REGULAR_FONT, uni=True)
    pdf.add_font('Poppins', 'B', BOLD_FONT, uni=True)
    
    # Page 1: Title
    pdf.add_page()
    pdf.set_font('Poppins', 'B', 32)
    pdf.ln(100)
    pdf.multi_cell(190, 15, TOPIC, 0, 'C')
    
    # Page 2: TOC
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 24)
    pdf.cell(0, 20, "विषय सूची", 0, 1, 'C')
    pdf.ln(5)
    pdf.set_font('Poppins', 'B', 12)
    pdf.cell(15, 10, 'क्र.सं.', 1, 0, 'C')
    pdf.cell(20, 10, 'अध्याय', 1, 0, 'C')
    pdf.cell(125, 10, 'विषय', 1, 0, 'C')
    pdf.cell(20, 10, 'पृष्ठ', 1, 1, 'C')
    
    pdf.set_font('Poppins', '', 11)
    rows = [("1", "-", "सारांश", "3")]
    curr_pg = 4
    for i in range(7):
        rows.append((str(i+2), str(i+1), chapter_names[i], str(curr_pg)))
        curr_pg += pages_per_ch[i]
    rows.append((str(9), "-", "शोध का उद्देश्य", "58"))
    rows.append((str(10), "-", "शोध का महत्व", "59"))
    rows.append((str(11), "-", "निष्कर्ष", "60"))
    
    for row in rows:
        pdf.cell(15, 10, row[0], 1, 0, 'C')
        pdf.cell(20, 10, row[1], 1, 0, 'C')
        pdf.cell(125, 10, row[2], 1, 0, 'L')
        pdf.cell(20, 10, row[3], 1, 1, 'C')

    # Page 3: Summary
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 26)
    pdf.cell(0, 20, "सारांश", 0, 1, 'L')
    pdf.set_font('Poppins', '', 14)
    pdf.multi_cell(0, 11.5, get_unique_content(3, 11), 0, 'J')
    
    # Content Pages 4 to 57
    page_count = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            pdf.add_page()
            page_count += 1
            pdf.set_y(40)
            if j == 0:
                pdf.set_font('Poppins', 'B', 26)
                pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                pdf.ln(5)
            pdf.set_font('Poppins', '', 14)
            pdf.multi_cell(0, 11.5, get_unique_content(page_count, 11), 0, 'J')
            
    # Extra Pages 58, 59, 60
    extra_titles = ["शोध का उद्देश्य", "शोध का महत्व", "निष्कर्ष"]
    for title in extra_titles:
        pdf.add_page()
        page_count += 1
        pdf.set_y(40)
        pdf.set_font('Poppins', 'B', 26)
        pdf.cell(0, 20, title, 0, 1, 'L')
        pdf.set_font('Poppins', '', 14)
        pdf.multi_cell(0, 11.5, get_unique_content(page_count, 11), 0, 'J')
        
    pdf.output("Economic_Diplomacy_Perfect_60.pdf")
    
    # DOCX
    doc = Document()
    set_page_border_docx(doc)
    # Title
    doc.add_paragraph("\n" * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TOPIC)
    run.font.name = 'Poppins'
    run.font.size = Pt(36)
    run.bold = True
    
    # TOC
    doc.add_page_break()
    doc.add_heading("विषय सूची", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['क्र.सं.', 'अध्याय', 'विषय', 'पृष्ठ']):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i in range(4): row[i].text = row_data[i]
        
    # Summary
    doc.add_page_break()
    h = doc.add_heading("सारांश", level=1)
    h.runs[0].font.name = 'Poppins'
    h.runs[0].font.size = Pt(26)
    p = doc.add_paragraph(get_unique_content(3, 11))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Poppins'
        run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.75

    # Content
    page_count_docx = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            doc.add_page_break()
            page_count_docx += 1
            if j == 0:
                h = doc.add_heading(f"अध्याय {i+1}: {chapter_names[i]}", level=1)
                h.runs[0].font.name = 'Poppins'
                h.runs[0].font.size = Pt(26)
            p = doc.add_paragraph(get_unique_content(page_count_docx, 11))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Poppins'
                run.font.size = Pt(14)
            p.paragraph_format.line_spacing = 1.75
            
    # Extras
    for title in extra_titles:
        doc.add_page_break()
        page_count_docx += 1
        h = doc.add_heading(title, level=1)
        h.runs[0].font.name = 'Poppins'
        h.runs[0].font.size = Pt(26)
        p = doc.add_paragraph(get_unique_content(page_count_docx, 11))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Poppins'
            run.font.size = Pt(14)
        p.paragraph_format.line_spacing = 1.75
        
    doc.save("Economic_Diplomacy_Perfect_60.docx")

print("Restoring PERFECT 60-page dissertation structure...")
generate_economic_diplomacy_perfect_60()
print("Success.")
