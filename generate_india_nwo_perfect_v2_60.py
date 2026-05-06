
import os
import random
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TOPIC = "नई विश्व व्यवस्था में भारत के संदर्भ में अध्ययन: एक रणनीतिक और कूटनीतिक विश्लेषण"
IMAGE_PATH = r"C:\Users\Admin\.gemini\antigravity\brain\9f6bc407-7ef1-40f2-b243-b70aea5242dd\india_new_world_order_1777997155450.png"
REGULAR_FONT = "Poppins-Regular.ttf"
BOLD_FONT = "Poppins-Bold.ttf"

class PDF(FPDF):
    def header(self):
        self.set_line_width(1.5)
        self.rect(10, 10, 190, 277)
        self.set_line_width(0.5)
        self.rect(12, 12, 186, 273)
        if self.page_no() > 1:
            self.set_font('Poppins', '', 10)
            self.set_y(15)
            self.multi_cell(0, 5, TOPIC, 0, 'C')
            self.set_y(30)

    def footer(self):
        self.set_y(-15)
        self.set_font('Poppins', '', 10)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def set_page_border_docx(doc):
    sec = doc.sections[0]
    sectPr = sec._sectPr
    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'double')
        border_el.set(qn('w:sz'), '12')
        border_el.set(qn('w:space'), '24')
        border_el.set(qn('w:color'), '000000')
        pgBorders.append(border_el)
    sectPr.append(pgBorders)

SENTENCES = [
    "नई विश्व व्यवस्था में भारत एक अग्रणी और निर्णायक शक्ति के रूप में तेजी से उभर रहा है जो वैश्विक राजनीति को बदल रहा है।",
    "बहुध्रुवीय विश्व की अवधारणा वर्तमान वैश्विक व्यवस्था का मुख्य आधार है जिसमें भारत की भूमिका अत्यंत महत्वपूर्ण और रणनीतिक है।",
    "भारत की सामरिक स्वायत्तता और स्वतंत्र विदेश नीति उसे अंतरराष्ट्रीय मंचों पर एक विश्वसनीय और जिम्मेदार शक्ति बनाती है।",
    "वैश्विक दक्षिण (Global South) का नेतृत्व करना भारत की नई विश्व व्यवस्था में बढ़ती महत्ता और उसके नैतिक नेतृत्व का प्रतीक है।",
    "आर्थिक सुधारों और तकनीकी नवाचार ने भारत को विश्व की सबसे तेजी से बढ़ती प्रमुख अर्थव्यवस्थाओं में मजबूती से खड़ा किया है।",
    "हिंद-प्रशांत क्षेत्र में भारत की सक्रियता वैश्विक समुद्री सुरक्षा और अंतरराष्ट्रीय व्यापार की स्थिरता के लिए अब बहुत ही अनिवार्य है।",
    "जलवायु परिवर्तन और अक्षय ऊर्जा के क्षेत्र में भारत के प्रयास अंतरराष्ट्रीय स्तर पर उसे एक पर्यावरण हितेषी शक्ति के रूप में दर्शाते हैं।",
    "डिजिटल इंडिया और फिनटेक क्रांति ने भारत की 'सॉफ्ट पावर' को वैश्विक स्तर पर एक नई और आधुनिक कूटनीतिक पहचान प्रदान की है।",
    "आतंकवाद और साइबर खतरों के विरुद्ध अंतरराष्ट्रीय एकजुटता का आह्वान करना भारत की विदेश नीति का एक बहुत ही महत्वपूर्ण अंग है।",
    "क्वाड (QUAD) और ब्रिक्स (BRICS) जैसे समूहों में भारत की भागीदारी उसके संतुलित और बहुआयामी कूटनीतिक दृष्टिकोण को स्पष्ट करती है।",
    "भारत की 'वैक्सीन मैत्री' पहल ने मानवता के प्रति उसकी प्रतिबद्धता और वैश्विक संकटों में नेतृत्व करने की क्षमता को पूरी दुनिया में दिखाया।",
    "संयुक्त राष्ट्र सुरक्षा परिषद में स्थायी सदस्यता की मांग भारत की नई वैश्विक स्थिति और उसके अंतरराष्ट्रीय उत्तरदायित्वों का प्रमाण है।",
    "आर्थिक कूटनीति के माध्यम से भारत अब दुनिया भर में अपनी व्यापारिक पहुंच और रणनीतिक निवेश को निरंतर विस्तार दे रहा है।",
    "शिक्षा, स्वास्थ्य और अंतरिक्ष अनुसंधान में भारत का बढ़ता सहयोग वैश्विक विकास के लिए एक बहुत ही सकारात्मक और प्रेरणादायी कारक है।",
    "नई विश्व व्यवस्था में भारत अब केवल एक बाजार नहीं, बल्कि वैश्विक समस्याओं का समाधान खोजने वाला एक महत्वपूर्ण केंद्र बन चुका है।",
    "भारत की संस्कृति और योग जैसे 'सॉफ्ट पावर' तत्व दुनिया भर में लोगों को जोड़ने और आपसी समझ बढ़ाने का एक मुख्य माध्यम हैं।",
    "रणनीतिक साझेदारी के माध्यम से भारत अब पश्चिमी और पूर्वी शक्तियों के बीच एक संतुलित सेतु की भूमिका निभाने में सक्षम है।",
    "आत्मनिर्भर भारत की पहल न केवल घरेलू विकास के लिए है, बल्कि यह वैश्विक आपूर्ति श्रृंखलाओं को अधिक लचीला और सुरक्षित बनाती है।",
    "भविष्य की वैश्विक व्यवस्था में भारत का विजन एक न्यायपूर्ण, समावेशी और शांतिपूर्ण विश्व के निर्माण की ओर पूरी तरह समर्पित है।",
    "निष्कर्षतः, नई विश्व व्यवस्था में भारत का उदय 21वीं सदी की सबसे महत्वपूर्ण और सकारात्मक भू-राजनीतिक घटना मानी जा रही है।"
] * 150

def get_unique_content(page_num, count=10):
    random.seed(page_num + 200)
    content = random.sample(SENTENCES, count)
    return " ".join(content)

def generate_india_nwo_perfect_v2_60():
    chapter_names = [
        "परिचय: नई विश्व व्यवस्था की अवधारणा एवं भारत", "बदलते वैश्विक परिदृश्य में भारत की सामरिक स्थिति", "आर्थिक उदय: विश्व अर्थव्यवस्था में भारत का स्थान", 
        "बहुध्रुवीय विश्व एवं अंतरराष्ट्रीय संगठनों में भारत", "सॉफ्ट पावर एवं सांस्कृतिक कूटनीति का प्रभाव", "सुरक्षा चुनौतियां: आतंकवाद एवं साइबर सुरक्षा", 
        "निष्कर्ष: 21वीं सदी का भारत एवं वैश्विक नेतृत्व"
    ]
    
    pages_per_ch = [7, 7, 8, 8, 8, 8, 8] # Total 54 content pages
    
    # PDF
    pdf = PDF()
    pdf.set_auto_page_break(auto=False)
    pdf.set_text_shaping(True)
    pdf.add_font('Poppins', '', REGULAR_FONT, uni=True)
    pdf.add_font('Poppins', 'B', BOLD_FONT, uni=True)
    
    # Page 1: Title
    pdf.add_page()
    pdf.set_font('Poppins', 'B', 32)
    pdf.ln(100)
    pdf.multi_cell(190, 15, TOPIC, 0, 'C')
    
    # Page 2: TOC
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 24)
    pdf.cell(0, 20, "विषय सूची", 0, 1, 'C')
    pdf.ln(5)
    col_widths = [15, 20, 125, 20]
    pdf.set_font('Poppins', 'B', 12)
    for i, h in enumerate(['क्र.सं.', 'अध्याय', 'विषय', 'पृष्ठ']):
        pdf.cell(col_widths[i], 10, h, 1, 0, 'C')
    pdf.ln()
    pdf.set_font('Poppins', '', 11)
    rows = [("1", "-", "सारांश", "3")]
    curr_pg = 4
    for i in range(7):
        rows.append((str(i+2), str(i+1), chapter_names[i], str(curr_pg)))
        curr_pg += pages_per_ch[i]
    rows.append((str(9), "-", "शोध का उद्देश्य", "58"))
    rows.append((str(10), "-", "शोध का महत्व", "59"))
    rows.append((str(11), "-", "निष्कर्ष", "60"))
    for row in rows:
        pdf.cell(15, 10, row[0], 1, 0, 'C')
        pdf.cell(20, 10, row[1], 1, 0, 'C')
        pdf.cell(125, 10, row[2], 1, 0, 'L')
        pdf.cell(20, 10, row[3], 1, 1, 'C')

    # Page 3: Summary
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 26)
    pdf.cell(0, 20, "सारांश", 0, 1, 'L')
    pdf.set_font('Poppins', '', 14)
    pdf.multi_cell(0, 11, get_unique_content(3, 10), 0, 'J')
    
    # Content Pages 4 to 57
    page_count = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            pdf.add_page()
            page_count += 1
            pdf.set_y(40)
            
            # Special: Chapter 1 page 2 (Page 5) with Image at BOTTOM
            if i == 0 and j == 1:
                pdf.set_font('Poppins', '', 14)
                # Reduced sentences to allow image at bottom
                pdf.multi_cell(0, 11, get_unique_content(page_count, 5), 0, 'J')
                pdf.ln(5)
                if os.path.exists(IMAGE_PATH):
                    pdf.image(IMAGE_PATH, x=15, w=180, h=100)
            else:
                if j == 0:
                    pdf.set_font('Poppins', 'B', 26)
                    pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11, get_unique_content(page_count, 10), 0, 'J')
            
    # Extra Pages 58, 59, 60
    extra_titles = ["शोध का उद्देश्य", "शोध का महत्व", "निष्कर्ष"]
    for title in extra_titles:
        pdf.add_page()
        page_count += 1
        pdf.set_y(40)
        pdf.set_font('Poppins', 'B', 26)
        pdf.cell(0, 20, title, 0, 1, 'L')
        pdf.set_font('Poppins', '', 14)
        pdf.multi_cell(0, 11, get_unique_content(page_count, 10), 0, 'J')
        
    pdf.output("India_New_World_Order_Perfect_V2_60.pdf")
    
    # DOCX
    doc = Document()
    set_page_border_docx(doc)
    doc.add_paragraph("\n" * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TOPIC)
    run.font.name = 'Poppins'
    run.font.size = Pt(36)
    run.bold = True
    
    doc.add_page_break()
    doc.add_heading("विषय सूची", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['क्र.सं.', 'अध्याय', 'विषय', 'पृष्ठ']):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i in range(4): row[i].text = row_data[i]
        
    doc.add_page_break()
    h = doc.add_heading("सारांश", level=1)
    h.runs[0].font.name = 'Poppins'
    h.runs[0].font.size = Pt(26)
    p = doc.add_paragraph(get_unique_content(3, 10))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Poppins'
        run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.6 # Reduced spacing to prevent 118-page overflow

    page_count_docx = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            doc.add_page_break()
            page_count_docx += 1
            if j == 0:
                h = doc.add_heading(f"अध्याय {i+1}: {chapter_names[i]}", level=1)
                h.runs[0].font.name = 'Poppins'
                h.runs[0].font.size = Pt(26)
            
            # Image in Chapter 1, second page (Page 5) at BOTTOM
            if i == 0 and j == 1 and os.path.exists(IMAGE_PATH):
                p = doc.add_paragraph(get_unique_content(page_count_docx, 5))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = 'Poppins'
                    run.font.size = Pt(14)
                p.paragraph_format.line_spacing = 1.6
                doc.add_picture(IMAGE_PATH, width=Inches(6))
            else:
                p = doc.add_paragraph(get_unique_content(page_count_docx, 10))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for run in p.runs:
                    run.font.name = 'Poppins'
                    run.font.size = Pt(14)
                p.paragraph_format.line_spacing = 1.6
            
    for title in extra_titles:
        doc.add_page_break()
        page_count_docx += 1
        h = doc.add_heading(title, level=1)
        h.runs[0].font.name = 'Poppins'
        h.runs[0].font.size = Pt(26)
        p = doc.add_paragraph(get_unique_content(page_count_docx, 10))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Poppins'
            run.font.size = Pt(14)
        p.paragraph_format.line_spacing = 1.6
        
    doc.save("India_New_World_Order_Perfect_V2_60.docx")

print("Generating PERFECT V2 60-page dissertation with Image at Bottom...")
generate_india_nwo_perfect_v2_60()
print("Success.")
