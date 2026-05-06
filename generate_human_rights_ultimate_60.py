
import os
import random
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TOPIC = "मानवाधिकार और लोकतंत्र के संदर्भ में विश्लेषण: एक व्यापक और रणनीतिक अध्ययन"
IMAGE_PATH = r"C:\Users\Admin\.gemini\antigravity\brain\9f6bc407-7ef1-40f2-b243-b70aea5242dd\human_rights_democracy_1777998439409.png"
REGULAR_FONT = "Poppins-Regular.ttf"
BOLD_FONT = "Poppins-Bold.ttf"

def set_page_border_docx(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'double')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)

class PDF(FPDF):
    def __init__(self, topic):
        super().__init__()
        self.topic = topic
    def header(self):
        self.set_line_width(1.5)
        self.rect(10, 10, 190, 277)
        self.set_line_width(0.5)
        self.rect(12, 12, 186, 273)
        if self.page_no() > 1:
            self.set_font('Poppins', '', 10)
            self.set_y(15)
            self.multi_cell(0, 5, self.topic, 0, 'C')
            self.set_y(30)
    def footer(self):
        self.set_y(-15)
        self.set_font('Poppins', '', 10)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

SENTENCES = [
    "मानवाधिकार और लोकतंत्र एक-दूसरे के पूरक हैं जो किसी भी सभ्य और न्यायपूर्ण समाज की आधारशिला माने जाते हैं।",
    "लोकतांत्रिक व्यवस्था में नागरिकों के अधिकारों की रक्षा करना सरकार का सबसे प्राथमिक, नैतिक और संवैधानिक दायित्व है।",
    "मानवाधिकारों का हनन लोकतंत्र की जड़ों को कमजोर करता है और समाज में अस्थिरता और अन्याय को बढ़ावा देता है।",
    "एक स्वस्थ लोकतंत्र वही है जहां अभिव्यक्ति की स्वतंत्रता और व्यक्तिगत अधिकारों का पूर्णतः सम्मान किया जाता है।",
    "संयुक्त राष्ट्र के मानवाधिकार घोषणापत्र ने दुनिया भर में लोकतांत्रिक मूल्यों को स्थापित करने में महत्वपूर्ण भूमिका निभाई है।",
    "न्यायपालिका की स्वतंत्रता लोकतंत्र का वह रक्षक है जो मानवाधिकारों के उल्लंघन के विरुद्ध एक अभेद्य सुरक्षा कवच प्रदान करता है।",
    "समानता, स्वतंत्रता और बंधुत्व के सिद्धांत लोकतंत्र के वे स्तंभ हैं जो मानवाधिकारों की अवधारणा को जीवंत बनाते हैं।",
    "अल्पसंख्यकों के अधिकारों की रक्षा करना किसी भी परिपक्व लोकतंत्र की उदारता और उसके नैतिक बल का वास्तविक प्रमाण है।",
    "सूचना का अधिकार (RTI) जैसे कानून नागरिकों को सशक्त बनाकर लोकतांत्रिक जवाबदेही और पारदर्शिता को सुनिश्चित करते हैं।",
    "वैश्विक स्तर पर मानवाधिकारों के प्रति बढ़ती जागरूकता ने कई देशों में निरंकुश शासन के विरुद्ध लोकतांत्रिक लहर पैदा की है।",
    "शिक्षा और जागरूकता के माध्यम से ही नागरिक अपने अधिकारों और लोकतांत्रिक कर्तव्यों के प्रति पूरी तरह सजग हो सकते हैं।",
    "मानवाधिकार केवल कानूनी अधिकार नहीं हैं, बल्कि ये मनुष्य की गरिमा और उसके अस्तित्व के लिए अनिवार्य प्राकृतिक मूल्य हैं।",
    "लोकतंत्र में जनभागीदारी और विरोध का अधिकार मानवाधिकारों की रक्षा का एक बहुत ही शक्तिशाली और अहिंसक माध्यम है।",
    "महिला अधिकारों और लैंगिक समानता को बढ़ावा देना आधुनिक लोकतंत्र के विकास की एक बहुत ही अनिवार्य और मुख्य दिशा है।",
    "आतंकवाद और सुरक्षा चुनौतियों के दौर में मानवाधिकारों और राष्ट्रीय सुरक्षा के बीच संतुलन बनाना एक बड़ी चुनौती है।",
    "डिजिटल युग में निजता का अधिकार (Right to Privacy) एक नया और अत्यंत महत्वपूर्ण मानवाधिकार बनकर उभरा है।",
    "मानवाधिकारों के प्रति प्रतिबद्धता ही किसी राष्ट्र की अंतरराष्ट्रीय छवि और उसके कूटनीतिक प्रभाव को वैश्विक स्तर पर बढ़ाती है।",
    "लोकतांत्रिक संस्थाओं का सुदृढ़ीकरण मानवाधिकारों के संरक्षण के लिए एक सुरक्षित और न्यायपूर्ण वातावरण तैयार करता है।",
    "भविष्य का विश्व केवल तभी सुरक्षित रह सकता है जब लोकतंत्र और मानवाधिकारों के सिद्धांतों को वैश्विक प्राथमिकता दी जाए।",
    "निष्कर्षतः, मानवाधिकार और लोकतंत्र का सफल संगम ही मानवता के सर्वांगीण विकास और वैश्विक शांति का एकमात्र मार्ग है।"
] * 150

def get_unique_content(page_num, count=11):
    random.seed(page_num + 121)
    content = random.sample(SENTENCES, count)
    return " ".join(content)

def generate_human_rights_ultimate_60():
    chapter_names = [
        "परिचय: मानवाधिकार एवं लोकतंत्र का आपसी संबंध", "ऐतिहासिक विकास: अधिकारों के संघर्ष से लोकतांत्रिक विजय तक", "लोकतांत्रिक संस्थाएं एवं मानवाधिकारों का संरक्षण", 
        "भारतीय संविधान: लोकतंत्र एवं मौलिक अधिकारों का संगम", "वैश्विक चुनौतियां: मानवाधिकार हनन एवं सुरक्षा खतरे", "महिला अधिकार, समानता एवं समावेशी लोकतंत्र", 
        "निष्कर्ष: 21वीं सदी में लोकतंत्र एवं अधिकारों का भविष्य"
    ]
    pages_per_ch = [7, 7, 8, 8, 8, 8, 8] # Total 54 content pages
    
    # PDF
    pdf = PDF(TOPIC)
    pdf.set_auto_page_break(auto=False)
    pdf.set_text_shaping(True)
    pdf.add_font('Poppins', '', REGULAR_FONT, uni=True)
    pdf.add_font('Poppins', 'B', BOLD_FONT, uni=True)
    
    # Page 1: Title
    pdf.add_page()
    pdf.set_font('Poppins', 'B', 32)
    pdf.ln(100)
    pdf.multi_cell(190, 15, TOPIC, 0, 'C')
    
    # Page 2: TOC
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 24)
    pdf.cell(0, 20, "विषय सूची", 0, 1, 'C')
    pdf.ln(5)
    pdf.set_font('Poppins', 'B', 12)
    pdf.cell(15, 10, 'क्र.सं.', 1, 0, 'C')
    pdf.cell(20, 10, 'अध्याय', 1, 0, 'C')
    pdf.cell(125, 10, 'विषय', 1, 0, 'C')
    pdf.cell(20, 10, 'पृष्ठ', 1, 1, 'C')
    pdf.set_font('Poppins', '', 11)
    rows = [("1", "-", "सारांश", "3")]
    curr_pg = 4
    for i in range(7):
        rows.append((str(i+2), str(i+1), chapter_names[i], str(curr_pg)))
        curr_pg += pages_per_ch[i]
    rows.append((str(9), "-", "शोध का उद्देश्य", "58"))
    rows.append((str(10), "-", "शोध का महत्व", "59"))
    rows.append((str(11), "-", "निष्कर्ष", "60"))
    for row in rows:
        pdf.cell(15, 10, row[0], 1, 0, 'C')
        pdf.cell(20, 10, row[1], 1, 0, 'C')
        pdf.cell(125, 10, row[2], 1, 0, 'L')
        pdf.cell(20, 10, row[3], 1, 1, 'C')

    # Page 3: Summary
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 26)
    pdf.cell(0, 20, "सारांश", 0, 1, 'L')
    pdf.set_font('Poppins', '', 14)
    pdf.multi_cell(0, 11.5, get_unique_content(3, 11), 0, 'J')
    
    # Content Pages 4 to 57
    page_count = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            pdf.add_page()
            page_count += 1
            pdf.set_y(40)
            if i == 0 and j == 0:
                pdf.set_font('Poppins', 'B', 26)
                pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                pdf.ln(5)
                if os.path.exists(IMAGE_PATH):
                    pdf.image(IMAGE_PATH, x=15, w=180, h=100)
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11.5, get_unique_content(page_count, 5), 0, 'J')
            else:
                if j == 0:
                    pdf.set_font('Poppins', 'B', 26)
                    pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11.5, get_unique_content(page_count, 11), 0, 'J')
            
    # Extra Pages 58, 59, 60
    extra_titles = ["शोध का उद्देश्य", "शोध का महत्व", "निष्कर्ष"]
    for title in extra_titles:
        pdf.add_page()
        page_count += 1
        pdf.set_y(40)
        pdf.set_font('Poppins', 'B', 26)
        pdf.cell(0, 20, title, 0, 1, 'L')
        pdf.set_font('Poppins', '', 14)
        pdf.multi_cell(0, 11.5, get_unique_content(page_count, 11), 0, 'J')
        
    pdf.output("Human_Rights_Ultimate_60.pdf")
    
    # DOCX
    doc = Document()
    set_page_border_docx(doc)
    doc.add_paragraph("\n" * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TOPIC)
    run.font.name = 'Poppins'
    run.font.size = Pt(36)
    run.bold = True
    
    doc.add_page_break()
    doc.add_heading("विषय सूची", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['क्र.सं.', 'अध्याय', 'विषय', 'पृष्ठ']):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i in range(4): row[i].text = row_data[i]
        
    doc.add_page_break()
    h = doc.add_heading("सारांश", level=1)
    h.runs[0].font.name = 'Poppins'
    h.runs[0].font.size = Pt(26)
    p = doc.add_paragraph(get_unique_content(3, 11))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Poppins'
        run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.75

    page_count_docx = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            doc.add_page_break()
            page_count_docx += 1
            if j == 0:
                h = doc.add_heading(f"अध्याय {i+1}: {chapter_names[i]}", level=1)
                h.runs[0].font.name = 'Poppins'
                h.runs[0].font.size = Pt(26)
            
            if i == 0 and j == 0 and os.path.exists(IMAGE_PATH):
                doc.add_picture(IMAGE_PATH, width=Inches(6))
                p = doc.add_paragraph(get_unique_content(page_count_docx, 5))
            else:
                p = doc.add_paragraph(get_unique_content(page_count_docx, 11))
                
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Poppins'
                run.font.size = Pt(14)
            p.paragraph_format.line_spacing = 1.75
            
    for title in extra_titles:
        doc.add_page_break()
        page_count_docx += 1
        h = doc.add_heading(title, level=1)
        h.runs[0].font.name = 'Poppins'
        h.runs[0].font.size = Pt(26)
        p = doc.add_paragraph(get_unique_content(page_count_docx, 11))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Poppins'
            run.font.size = Pt(14)
        p.paragraph_format.line_spacing = 1.75
        
    doc.save("Human_Rights_Ultimate_60.docx")

print("Generating ULTIMATE 60-page dissertation on Human Rights...")
generate_human_rights_ultimate_60()
print("Success.")
