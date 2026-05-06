
import os
import random
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --- SHARED CONFIG ---
REGULAR_FONT = "Poppins-Regular.ttf"
BOLD_FONT = "Poppins-Bold.ttf"

def set_page_border_docx(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'double')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)

class PDF(FPDF):
    def __init__(self, topic):
        super().__init__()
        self.topic = topic
    def header(self):
        self.set_line_width(1.5)
        self.rect(10, 10, 190, 277)
        self.set_line_width(0.5)
        self.rect(12, 12, 186, 273)
        if self.page_no() > 1:
            self.set_font('Poppins', '', 10)
            self.set_y(15)
            self.multi_cell(0, 5, self.topic, 0, 'C')
            self.set_y(30)
    def footer(self):
        self.set_y(-15)
        self.set_font('Poppins', '', 10)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def get_unique_content(sentences, page_num, seed_val, count=11):
    random.seed(page_num + seed_val)
    content = random.sample(sentences, count)
    return " ".join(content)

# --- NWO CONTENT ---
NWO_TOPIC = "नई विश्व व्यवस्था में भारत के संदर्भ में अध्ययन: एक रणनीतिक और कूटनीतिक विश्लेषण"
NWO_IMAGE = r"C:\Users\Admin\.gemini\antigravity\brain\9f6bc407-7ef1-40f2-b243-b70aea5242dd\india_new_world_order_1777997155450.png"
NWO_SENTENCES = [
    "नई विश्व व्यवस्था में भारत एक अग्रणी और निर्णायक शक्ति के रूप में तेजी से उभर रहा है जो वैश्विक राजनीति को बदल रहा है।",
    "बहुध्रुवीय विश्व की अवधारणा वर्तमान वैश्विक व्यवस्था का मुख्य आधार है जिसमें भारत की भूमिका अत्यंत महत्वपूर्ण और रणनीतिक है।",
    "भारत की सामरिक स्वायत्तता और स्वतंत्र विदेश नीति उसे अंतरराष्ट्रीय मंचों पर एक विश्वसनीय और जिम्मेदार शक्ति बनाती है।",
    "वैश्विक दक्षिण (Global South) का नेतृत्व करना भारत की नई विश्व व्यवस्था में बढ़ती महत्ता और उसके नैतिक नेतृत्व का प्रतीक है।",
    "आर्थिक सुधारों और तकनीकी नवाचार ने भारत को विश्व की सबसे तेजी से बढ़ती प्रमुख अर्थव्यवस्थाओं में मजबूती से खड़ा किया है।",
    "हिंद-प्रशांत क्षेत्र में भारत की सक्रियता वैश्विक समुद्री सुरक्षा और अंतरराष्ट्रीय व्यापार की स्थिरता के लिए अब बहुत ही अनिवार्य है।",
    "जलवायु परिवर्तन और अक्षय ऊर्जा के क्षेत्र में भारत के प्रयास अंतरराष्ट्रीय स्तर पर उसे एक पर्यावरण हितेषी शक्ति के रूप में दर्शाते हैं।",
    "डिजिटल इंडिया और फिनटेक क्रांति ने भारत की 'सॉफ्ट पावर' को वैश्विक स्तर पर एक नई और आधुनिक कूटनीतिक पहचान प्रदान की है।",
    "आतंकवाद और साइबर खतरों के विरुद्ध अंतरराष्ट्रीय एकजुटता का आह्वान करना भारत की विदेश नीति का एक बहुत ही महत्वपूर्ण अंग है।",
    "क्वाड (QUAD) और ब्रिक्स (BRICS) जैसे समूहों में भारत की भागीदारी उसके संतुलित और बहुआयामी कूटनीतिक दृष्टिकोण को स्पष्ट करती है।",
    "भारत की 'वैक्सीन मैत्री' पहल ने मानवता के प्रति उसकी प्रतिबद्धता और वैश्विक संकटों में नेतृत्व करने की क्षमता को पूरी दुनिया में दिखाया।",
    "संयुक्त राष्ट्र सुरक्षा परिषद में स्थायी सदस्यता की मांग भारत की नई वैश्विक स्थिति और उसके अंतरराष्ट्रीय उत्तरदायित्वों का प्रमाण है।",
    "आर्थिक कूटनीति के माध्यम से भारत अब दुनिया भर में अपनी व्यापारिक पहुंच और रणनीतिक निवेश को निरंतर विस्तार दे रहा है।",
    "शिक्षा, स्वास्थ्य और अंतरिक्ष अनुसंधान में भारत का बढ़ता सहयोग वैश्विक विकास के लिए एक बहुत ही सकारात्मक और प्रेरणादायी कारक है।",
    "नई विश्व व्यवस्था में भारत अब केवल एक बाजार नहीं, बल्कि वैश्विक समस्याओं का समाधान खोजने वाला एक महत्वपूर्ण केंद्र बन चुका है।",
    "भारत की संस्कृति और योग जैसे 'सॉफ्ट पावर' तत्व दुनिया भर में लोगों को जोड़ने और आपसी समझ बढ़ाने का एक मुख्य माध्यम हैं।",
    "रणनीतिक साझेदारी के माध्यम से भारत अब पश्चिमी और पूर्वी शक्तियों के बीच एक संतुलित सेतु की भूमिका निभाने में सक्षम है।",
    "आत्मनिर्भर भारत की पहल न केवल घरेलू विकास के लिए है, बल्कि यह वैश्विक आपूर्ति श्रृंखलाओं को अधिक लचीला और सुरक्षित बनाती है।",
    "भविष्य की वैश्विक व्यवस्था में भारत का विजन एक न्यायपूर्ण, समावेशी और शांतिपूर्ण विश्व के निर्माण की ओर पूरी तरह समर्पित है।",
    "निष्कर्षतः, नई विश्व व्यवस्था में भारत का उदय 21वीं सदी की सबसे महत्वपूर्ण और सकारात्मक भू-राजनीतिक घटना मानी जा रही है।"
] * 150

# --- ECONOMIC CONTENT ---
ECO_TOPIC = "आर्थिक विकास और कूटनीति के संदर्भ में विश्लेषण: एक रणनीतिक अध्ययन"
ECO_SENTENCES = [
    "आर्थिक कूटनीति वर्तमान वैश्विक व्यवस्था में किसी भी राष्ट्र की विदेश नीति का एक अत्यंत महत्वपूर्ण और आधारभूत स्तंभ बन गई है।",
    "आर्थिक विकास और कूटनीति का आपसी संबंध राष्ट्रों के बीच शक्ति संतुलन और सहयोग के नए आयामों को निरंतर परिभाषित कर रहा है।",
    "वैश्वीकरण के इस युग में व्यापारिक समझौते और विदेशी निवेश को आकर्षित करना कूटनीतिक सफलता का एक प्रमुख पैमाना माना जाता है।",
    "आर्थिक विकास केवल घरेलू नीतियों पर निर्भर नहीं है, बल्कि यह अंतरराष्ट्रीय व्यापारिक संबंधों और कूटनीतिक संवाद पर भी टिका है।",
    "भारत जैसी उभरती अर्थव्यवस्थाओं के लिए आर्थिक कूटनीति वैश्विक स्तर पर अपनी रणनीतिक स्वायत्तता को मजबूत करने का एक साधन है।",
    "ऊर्जा सुरक्षा और संसाधनों तक पहुंच सुनिश्चित करना आर्थिक कूटनीति का एक बहुत ही अनिवार्य, चुनौतीपूर्ण और रणनीतिक पक्ष है।",
    "विदेशी प्रत्यक्ष निवेश (FDI) को बढ़ावा देने के लिए राष्ट्रों के बीच स्वस्थ कूटनीतिक संबंधों का होना आज के दौर में बहुत जरूरी है।",
    "आर्थिक कूटनीति के माध्यम से राष्ट्र अपने आर्थिक हितों की रक्षा करते हुए वैश्विक राजनीति में अपना प्रभाव और स्थान बढ़ाते हैं।",
    "बहुपक्षीय संस्थानों जैसे विश्व व्यापार संगठन (WTO) में अपनी आवाज बुलंद करना विकासशील देशों की आर्थिक कूटनीति का हिस्सा है।",
    "आर्थिक विकास के लिए तकनीकी हस्तांतरण और नवाचार को बढ़ावा देना कूटनीतिक समझौतों का एक बहुत ही महत्वपूर्ण और आधुनिक लक्ष्य है।",
    "आर्थिक कूटनीति न केवल व्यापार को बढ़ावा देती है, बल्कि यह राष्ट्रों के बीच आपसी विश्वास और दीर्घकालिक शांति का आधार भी बनती है।",
    "क्षेत्रीय व्यापारिक ब्लॉक (Regional Trading Blocks) आर्थिक कूटनीति के माध्यम से क्षेत्रीय आर्थिक एकीकरण को मजबूती प्रदान करते हैं।",
    "सतत विकास लक्ष्यों (SDGs) को प्राप्त करने के लिए अंतरराष्ट्रीय आर्थिक सहयोग और कूटनीतिक समन्वय आज अत्यंत अनिवार्य हो गया है।",
    "आर्थिक प्रतिबंधों का कूटनीतिक हथियार के रूप में उपयोग वैश्विक व्यापार और अंतरराष्ट्रीय संबंधों की स्थिरता के लिए एक बड़ी चुनौती है।",
    "डिजिटल अर्थव्यवस्था और साइबर व्यापार के बढ़ते प्रभाव ने आर्थिक कूटनीति के लिए नए नियम और रणनीतिक ढांचे तैयार किए हैं।",
    "आर्थिक कूटनीति का मुख्य उद्देश्य राष्ट्र के आर्थिक विकास को गति देना और वैश्विक संसाधनों का न्यायोचित वितरण सुनिश्चित करना है।",
    "विकासशील राष्ट्रों के बीच दक्षिण-दक्षिण सहयोग (South-South Cooperation) आर्थिक कूटनीति का एक बहुत ही सकारात्मक स्वरूप है।",
    "आर्थिक कूटनीति और राष्ट्रीय सुरक्षा एक-दूसरे के पूरक हैं, जो राष्ट्र को वैश्विक स्तर पर अधिक सुरक्षित और समृद्ध बनाते हैं।",
    "पर्यावरण अनुकूल व्यापार नीतियां और हरित ऊर्जा के क्षेत्र में सहयोग आधुनिक आर्थिक कूटनीति की एक नई और अनिवार्य दिशा है।",
    "निष्कर्षतः, आर्थिक विकास और कूटनीति का सफल संगम ही 21वीं सदी में एक सशक्त और समृद्ध राष्ट्र निर्माण की सच्ची गारंटी है।"
] * 150

def generate_dissertation(topic, sentences, image_path, filename_base, seed_val, extra_sections=True):
    chapter_names = [
        "परिचय एवं अवधारणात्मक ढांचा", "ऐतिहासिक परिप्रेक्ष्य एवं विकास", "प्रमुख वैश्विक चुनौतियां एवं कूटनीति", 
        "भारत की भूमिका: रणनीतिक एवं आर्थिक विश्लेषण", "क्षेत्रीय एवं वैश्विक सहयोग के नए आयाम", "भविष्य की संभावनाएं एवं सुरक्षा खतरे", 
        "निष्कर्ष एवं नीतिगत सुझाव"
    ]
    pages_per_ch = [7, 7, 8, 8, 8, 8, 8]
    
    # DOCX
    doc = Document()
    set_page_border_docx(doc)
    doc.add_paragraph("\n" * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(topic)
    run.font.name = 'Poppins'
    run.font.size = Pt(36)
    run.bold = True
    
    doc.add_page_break()
    doc.add_heading("विषय सूची", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['क्र.सं.', 'अध्याय', 'विषय', 'पृष्ठ']):
        table.rows[0].cells[i].text = h
    
    rows = [("1", "-", "सारांश", "3")]
    curr_pg = 4
    for i in range(7):
        rows.append((str(i+2), str(i+1), chapter_names[i], str(curr_pg)))
        curr_pg += pages_per_ch[i]
    if extra_sections:
        rows.append((str(9), "-", "शोध का उद्देश्य", "58"))
        rows.append((str(10), "-", "शोध का महत्व", "59"))
        rows.append((str(11), "-", "निष्कर्ष", "60"))
    
    for row_data in rows:
        row = table.add_row().cells
        for i in range(4): row[i].text = row_data[i]
        
    doc.add_page_break()
    h = doc.add_heading("सारांश", level=1)
    h.runs[0].font.name = 'Poppins'
    h.runs[0].font.size = Pt(26)
    p = doc.add_paragraph(get_unique_content(sentences, 3, seed_val))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Poppins'
        run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.75

    page_count = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            doc.add_page_break()
            page_count += 1
            if j == 0:
                h = doc.add_heading(f"अध्याय {i+1}: {chapter_names[i]}", level=1)
                h.runs[0].font.name = 'Poppins'
                h.runs[0].font.size = Pt(26)
            
            if i == 0 and j == 0 and image_path and os.path.exists(image_path):
                doc.add_picture(image_path, width=Inches(6))
                p = doc.add_paragraph(get_unique_content(sentences, page_count, seed_val, 5))
            else:
                p = doc.add_paragraph(get_unique_content(sentences, page_count, seed_val))
                
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Poppins'
                run.font.size = Pt(14)
            p.paragraph_format.line_spacing = 1.75
            
    if extra_sections:
        for title in ["शोध का उद्देश्य", "शोध का महत्व", "निष्कर्ष"]:
            doc.add_page_break()
            page_count += 1
            h = doc.add_heading(title, level=1)
            h.runs[0].font.name = 'Poppins'
            h.runs[0].font.size = Pt(26)
            p = doc.add_paragraph(get_unique_content(sentences, page_count, seed_val))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Poppins'
                run.font.size = Pt(14)
            p.paragraph_format.line_spacing = 1.75
            
    doc.save(f"{filename_base}.docx")
    
    # PDF
    pdf = PDF(topic)
    pdf.set_auto_page_break(auto=False)
    pdf.set_text_shaping(True)
    pdf.add_font('Poppins', '', REGULAR_FONT, uni=True)
    pdf.add_font('Poppins', 'B', BOLD_FONT, uni=True)
    
    pdf.add_page()
    pdf.set_font('Poppins', 'B', 32)
    pdf.ln(100)
    pdf.multi_cell(190, 15, topic, 0, 'C')
    
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 24)
    pdf.cell(0, 20, "विषय सूची", 0, 1, 'C')
    pdf.ln(5)
    pdf.set_font('Poppins', 'B', 12)
    pdf.cell(15, 10, 'क्र.सं.', 1, 0, 'C')
    pdf.cell(20, 10, 'अध्याय', 1, 0, 'C')
    pdf.cell(125, 10, 'विषय', 1, 0, 'C')
    pdf.cell(20, 10, 'पृष्ठ', 1, 1, 'C')
    pdf.set_font('Poppins', '', 11)
    for row in rows:
        pdf.cell(15, 10, row[0], 1, 0, 'C')
        pdf.cell(20, 10, row[1], 1, 0, 'C')
        pdf.cell(125, 10, row[2], 1, 0, 'L')
        pdf.cell(20, 10, row[3], 1, 1, 'C')

    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 26)
    pdf.cell(0, 20, "सारांश", 0, 1, 'L')
    pdf.set_font('Poppins', '', 14)
    pdf.multi_cell(0, 11.5, get_unique_content(sentences, 3, seed_val), 0, 'J')
    
    page_count = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            pdf.add_page()
            page_count += 1
            pdf.set_y(40)
            if i == 0 and j == 0:
                pdf.set_font('Poppins', 'B', 26)
                pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                pdf.ln(5)
                if image_path and os.path.exists(image_path):
                    pdf.image(image_path, x=15, w=180, h=100)
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11.5, get_unique_content(sentences, page_count, seed_val, 5), 0, 'J')
            else:
                if j == 0:
                    pdf.set_font('Poppins', 'B', 26)
                    pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11.5, get_unique_content(sentences, page_count, seed_val), 0, 'J')
                
    if extra_sections:
        for title in ["शोध का उद्देश्य", "शोध का महत्व", "निष्कर्ष"]:
            pdf.add_page()
            page_count += 1
            pdf.set_y(40)
            pdf.set_font('Poppins', 'B', 26)
            pdf.cell(0, 20, title, 0, 1, 'L')
            pdf.set_font('Poppins', '', 14)
            pdf.multi_cell(0, 11.5, get_unique_content(sentences, page_count, seed_val), 0, 'J')
            
    pdf.output(f"{filename_base}.pdf")

print("Generating ULTIMATE versions with guaranteed borders...")
generate_dissertation(NWO_TOPIC, NWO_SENTENCES, NWO_IMAGE, "India_NWO_Ultimate_60", 200)
generate_dissertation(ECO_TOPIC, ECO_SENTENCES, None, "Economic_Diplomacy_Ultimate_60", 119)
print("Success.")
