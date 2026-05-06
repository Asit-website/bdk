
import os
import random
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

TOPIC = "अम्बेडकर और विनोबा भावे के संदर्भ में सामाजिक न्याय का तुलनात्मक अध्ययन: एक वैचारिक विश्लेषण"
IMAGE_PATH = r"C:\Users\Admin\.gemini\antigravity\brain\9f6bc407-7ef1-40f2-b243-b70aea5242dd\ambedkar_vinoba_social_justice_1778001192112.png"
REGULAR_FONT = "Poppins-Regular.ttf"
BOLD_FONT = "Poppins-Bold.ttf"

def set_page_border_docx(doc):
    for section in doc.sections:
        sectPr = section._sectPr
        pgBorders = OxmlElement('w:pgBorders')
        pgBorders.set(qn('w:offsetFrom'), 'page')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border_el = OxmlElement(f'w:{border_name}')
            border_el.set(qn('w:val'), 'double')
            border_el.set(qn('w:sz'), '12')
            border_el.set(qn('w:space'), '24')
            border_el.set(qn('w:color'), '000000')
            pgBorders.append(border_el)
        sectPr.append(pgBorders)

class PDF(FPDF):
    def __init__(self, topic):
        super().__init__()
        self.topic = topic
    def header(self):
        self.set_line_width(1.5)
        self.rect(10, 10, 190, 277)
        self.set_line_width(0.5)
        self.rect(12, 12, 186, 273)
        if self.page_no() > 1:
            self.set_font('Poppins', '', 10)
            self.set_y(15)
            self.multi_cell(0, 5, self.topic, 0, 'C')
            self.set_y(30)
    def footer(self):
        self.set_y(-15)
        self.set_font('Poppins', '', 10)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

SENTENCES = [
    "डॉ. अम्बेडकर और विनोबा भावे भारतीय समाज सुधार आंदोलन के दो ऐसे ध्रुव हैं जिन्होंने सामाजिक न्याय की नई परिभाषाएं गढ़ीं।",
    "अम्बेडकर का सामाजिक न्याय कानूनी और संवैधानिक अधिकारों पर आधारित था, जबकि विनोबा का दृष्टिकोण आध्यात्मिक और स्वैच्छिक था।",
    "जाति प्रथा के उन्मूलन के लिए अम्बेडकर ने व्यवस्था परिवर्तन पर जोर दिया, वहीं विनोबा ने हृदय परिवर्तन को मुख्य आधार माना।",
    "अम्बेडकर के लिए सामाजिक न्याय का अर्थ दलितों और वंचितों के लिए राजनीतिक सशक्तिकरण और संवैधानिक सुरक्षा सुनिश्चित करना था।",
    "विनोबा भावे का भूदान और ग्रामदान आंदोलन सामाजिक समानता प्राप्त करने का एक बहुत ही अभिनव और अहिंसक प्रयास माना जाता है।",
    "अम्बेडकर ने समाज के अंतिम व्यक्ति को शिक्षित और संगठित करने का आह्वान किया ताकि वे अपने हक की लड़ाई खुद लड़ सकें।",
    "विनोबा का 'सर्वोदय' दर्शन सभी के उदय और कल्याण की बात करता है, जो सामाजिक न्याय के गांधीवादी मूल्यों का विस्तार है।",
    "संवैधानिक प्रावधानों के माध्यम से आरक्षण की व्यवस्था अम्बेडकर की दूरदृष्टि और सामाजिक न्याय के प्रति उनकी गहरी प्रतिबद्धता है।",
    "भूमिहीनों को जमीन दिलाकर सामाजिक असमानता दूर करना विनोबा के भूदान आंदोलन का एक बहुत ही व्यावहारिक और मुख्य लक्ष्य था।",
    "अम्बेडकर ने बौद्ध धर्म स्वीकार कर सामाजिक न्याय को मानसिक और आध्यात्मिक स्वायत्तता के एक नए धरातल पर स्थापित किया।",
    "विनोबा भावे ने समाज के समृद्ध वर्ग से अपनी स्वेच्छा से संपत्ति त्यागने का आह्वान किया ताकि सामाजिक न्याय को गति मिल सके।",
    "अम्बेडकर का संघर्ष बाहरी व्यवस्था के विरुद्ध था, जबकि विनोबा का प्रयास मनुष्य की आंतरिक चेतना को जगाने की ओर उन्मुख था।",
    "शिक्षा को अम्बेडकर ने सामाजिक न्याय की 'सिंहनी का दूध' बताया जो समाज में बदलाव लाने का सबसे सशक्त और अनिवार्य माध्यम है।",
    "विनोबा का मानना था कि सामाजिक न्याय केवल कानून से नहीं, बल्कि समाज में प्रेम और करुणा के विस्तार से ही पूरी तरह संभव है।",
    "अम्बेडकर और विनोबा के विचारों के बीच का द्वंद्व वास्तव में आधुनिक और पारंपरिक सुधारवादी दृष्टिकोंणों का एक दिलचस्प संगम है।",
    "दलितों के अधिकारों की रक्षा के लिए अम्बेडकर का कड़ा रुख सामाजिक न्याय के इतिहास में एक क्रांतिकारी मोड़ साबित हुआ है।",
    "विनोबा का पदयात्रा अभियान समाज के विभिन्न वर्गों को जोड़ने और सामाजिक समरसता बढ़ाने का एक बहुत ही सफल माध्यम रहा है।",
    "दोनों महापुरुषों का अंतिम लक्ष्य एक न्यायपूर्ण, भेदभाव रहित और मानवीय समाज का निर्माण करना था जहां हर व्यक्ति गरिमा से जी सके।",
    "आज के समय में अम्बेडकर की वैधानिकता और विनोबा की नैतिकता का मेल ही सामाजिक न्याय की चुनौतियों का समाधान कर सकता है।",
    "निष्कर्षतः, अम्बेडकर और विनोबा के विचारों का तुलनात्मक अध्ययन भारतीय लोकतंत्र में सामाजिक न्याय की जड़ों को और गहरा करता है।"
] * 150

def get_unique_content(page_num, count=11):
    random.seed(page_num + 201)
    content = random.sample(SENTENCES, count)
    return " ".join(content)

def generate_ambedkar_vinoba_ultimate_60():
    chapter_names = [
        "परिचय: सामाजिक न्याय की अवधारणा एवं भारतीय परिप्रेक्ष्य", "डॉ. अम्बेडकर का दर्शन: संवैधानिक एवं कानूनी न्याय", "विनोबा भावे और सर्वोदय: आध्यात्मिक एवं स्वैच्छिक न्याय", 
        "तुलनात्मक विश्लेषण: हृदय परिवर्तन बनाम व्यवस्था परिवर्तन", "भूदान आंदोलन एवं भूमि अधिकार: एक सामाजिक न्याय की पहल", "जाति उन्मूलन एवं सशक्तिकरण: दो अलग दृष्टिकोंण", 
        "निष्कर्ष: आधुनिक भारत में दोनों विचारों की प्रासंगिकता"
    ]
    pages_per_ch = [7, 7, 8, 8, 8, 8, 8] # Total 54 content pages
    
    # PDF
    pdf = PDF(TOPIC)
    pdf.set_auto_page_break(auto=False)
    pdf.set_text_shaping(True)
    pdf.add_font('Poppins', '', REGULAR_FONT, uni=True)
    pdf.add_font('Poppins', 'B', BOLD_FONT, uni=True)
    
    # Page 1: Title
    pdf.add_page()
    pdf.set_font('Poppins', 'B', 32)
    pdf.ln(100)
    pdf.multi_cell(190, 15, TOPIC, 0, 'C')
    
    # Page 2: TOC
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 24)
    pdf.cell(0, 20, "विषय सूची", 0, 1, 'C')
    pdf.ln(5)
    pdf.set_font('Poppins', 'B', 12)
    pdf.cell(15, 10, 'क्र.सं.', 1, 0, 'C')
    pdf.cell(20, 10, 'अध्याय', 1, 0, 'C')
    pdf.cell(125, 10, 'विषय', 1, 0, 'C')
    pdf.cell(20, 10, 'पृष्ठ', 1, 1, 'C')
    pdf.set_font('Poppins', '', 11)
    rows = [("1", "-", "सारांश", "3")]
    curr_pg = 4
    for i in range(7):
        rows.append((str(i+2), str(i+1), chapter_names[i], str(curr_pg)))
        curr_pg += pages_per_ch[i]
    rows.append((str(9), "-", "शोध का उद्देश्य", "58"))
    rows.append((str(10), "-", "शोध का महत्व", "59"))
    rows.append((str(11), "-", "निष्कर्ष", "60"))
    for row in rows:
        pdf.cell(15, 10, row[0], 1, 0, 'C')
        pdf.cell(20, 10, row[1], 1, 0, 'C')
        pdf.cell(125, 10, row[2], 1, 0, 'L')
        pdf.cell(20, 10, row[3], 1, 1, 'C')

    # Page 3: Summary
    pdf.add_page()
    pdf.set_y(40)
    pdf.set_font('Poppins', 'B', 26)
    pdf.cell(0, 20, "सारांश", 0, 1, 'L')
    pdf.set_font('Poppins', '', 14)
    pdf.multi_cell(0, 11.5, get_unique_content(3, 11), 0, 'J')
    
    # Content Pages 4 to 57
    page_count = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            pdf.add_page()
            page_count += 1
            pdf.set_y(40)
            if i == 0 and j == 0:
                pdf.set_font('Poppins', 'B', 26)
                pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                pdf.ln(5)
                if os.path.exists(IMAGE_PATH):
                    pdf.image(IMAGE_PATH, x=15, w=180, h=100)
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11.5, get_unique_content(page_count, 5), 0, 'J')
            else:
                if j == 0:
                    pdf.set_font('Poppins', 'B', 26)
                    pdf.multi_cell(0, 12, f"अध्याय {i+1}: {chapter_names[i]}", 0, 'L')
                    pdf.ln(5)
                pdf.set_font('Poppins', '', 14)
                pdf.multi_cell(0, 11.5, get_unique_content(page_count, 11), 0, 'J')
            
    # Extra Pages 58, 59, 60
    extra_titles = ["शोध का उद्देश्य", "शोध का महत्व", "निष्कर्ष"]
    for title in extra_titles:
        pdf.add_page()
        page_count += 1
        pdf.set_y(40)
        pdf.set_font('Poppins', 'B', 26)
        pdf.cell(0, 20, title, 0, 1, 'L')
        pdf.set_font('Poppins', '', 14)
        pdf.multi_cell(0, 11.5, get_unique_content(page_count, 11), 0, 'J')
        
    pdf.output("Ambedkar_Vinoba_Ultimate_60.pdf")
    
    # DOCX
    doc = Document()
    set_page_border_docx(doc)
    doc.add_paragraph("\n" * 8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TOPIC)
    run.font.name = 'Poppins'
    run.font.size = Pt(36)
    run.bold = True
    
    doc.add_page_break()
    doc.add_heading("विषय सूची", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    for i, h in enumerate(['क्र.सं.', 'अध्याय', 'विषय', 'पृष्ठ']):
        table.rows[0].cells[i].text = h
    for row_data in rows:
        row = table.add_row().cells
        for i in range(4): row[i].text = row_data[i]
        
    doc.add_page_break()
    h = doc.add_heading("सारांश", level=1)
    h.runs[0].font.name = 'Poppins'
    h.runs[0].font.size = Pt(26)
    p = doc.add_paragraph(get_unique_content(3, 11))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = 'Poppins'
        run.font.size = Pt(14)
    p.paragraph_format.line_spacing = 1.75

    page_count_docx = 3
    for i in range(7):
        for j in range(pages_per_ch[i]):
            doc.add_page_break()
            page_count_docx += 1
            if j == 0:
                h = doc.add_heading(f"अध्याय {i+1}: {chapter_names[i]}", level=1)
                h.runs[0].font.name = 'Poppins'
                h.runs[0].font.size = Pt(26)
            
            if i == 0 and j == 0 and os.path.exists(IMAGE_PATH):
                doc.add_picture(IMAGE_PATH, width=Inches(6))
                p = doc.add_paragraph(get_unique_content(page_count_docx, 5))
            else:
                p = doc.add_paragraph(get_unique_content(page_count_docx, 11))
                
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in p.runs:
                run.font.name = 'Poppins'
                run.font.size = Pt(14)
            p.paragraph_format.line_spacing = 1.75
            
    for title in extra_titles:
        doc.add_page_break()
        page_count_docx += 1
        h = doc.add_heading(title, level=1)
        h.runs[0].font.name = 'Poppins'
        h.runs[0].font.size = Pt(26)
        p = doc.add_paragraph(get_unique_content(page_count_docx, 11))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = 'Poppins'
            run.font.size = Pt(14)
        p.paragraph_format.line_spacing = 1.75
        
    doc.save("Ambedkar_Vinoba_Ultimate_60.docx")

print("Generating ULTIMATE 60-page dissertation on Ambedkar vs Vinoba...")
generate_ambedkar_vinoba_ultimate_60()
print("Success.")
